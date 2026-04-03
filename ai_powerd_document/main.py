import os, base64, json, tempfile, re
from fastapi import FastAPI, HTTPException, Header
from fastapi.responses import HTMLResponse
from pydantic import BaseModel
from google import genai
import fitz
from docx import Document
import pytesseract
from PIL import Image
from dotenv import load_dotenv

load_dotenv()

app = FastAPI(title="DocuLens AI", version="3.0.0")
API_KEY = os.getenv("API_KEY", "sk_track2_987654321")
client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))

class DocumentRequest(BaseModel):
    fileName: str
    fileType: str
    fileBase64: str

class DocumentResponse(BaseModel):
    status: str
    fileName: str
    summary: str
    entities: dict
    sentiment: str

def extract_pdf(data):
    pages = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            pages.append(page.get_text("text"))
    return "\n".join(pages).strip()

def extract_docx(data):
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        f.write(data); path = f.name
    try:
        doc = Document(path)
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                r = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if r: lines.append(r)
        return "\n".join(lines)
    finally:
        os.unlink(path)

def extract_image(data):
    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as f:
        f.write(data); path = f.name
    try:
        img = Image.open(path)
        best = ""
        for cfg in ["--oem 3 --psm 6", "--oem 3 --psm 3", "--oem 3 --psm 11"]:
            r = pytesseract.image_to_string(img, config=cfg).strip()
            if len(r) > len(best): best = r
        return best
    finally:
        os.unlink(path)

def get_text(file_type, data):
    ft = file_type.lower().strip()
    if ft == "pdf": return extract_pdf(data)
    elif ft in ("docx","doc"): return extract_docx(data)
    elif ft in ("image","jpg","jpeg","png","tiff","bmp","webp","gif"): return extract_image(data)
    else: raise ValueError(f"Unsupported fileType: '{file_type}'")

PROMPT = """Analyze this document and return ONLY a valid JSON object. No markdown, no explanation.

Document: {filename}
Content:
{content}

Return exactly this structure:
{{
  "summary": "2-3 sentence factual summary",
  "entities": {{
    "names": ["person names"],
    "dates": ["dates found"],
    "organizations": ["org names"],
    "amounts": ["monetary or numeric values with units"]
  }},
  "sentiment": "Neutral"
}}

Rules: entities use [] if none found. sentiment is exactly Positive, Neutral, or Negative. Return ONLY JSON."""

def analyze(text, filename):
    content = text[:6000]
    resp = client.models.generate_content(model="gemini-1.5-flash", contents=PROMPT.format(filename=filename, content=content))
    raw = resp.text.strip()
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw).strip()
    try:
        result = json.loads(raw)
    except:
        m = re.search(r"\{.*\}", raw, re.DOTALL)
        result = json.loads(m.group()) if m else {}
    e = result.get("entities", {})
    return {
        "summary": str(result.get("summary","")).strip(),
        "entities": {
            "names": e.get("names",[]),
            "dates": e.get("dates",[]),
            "organizations": e.get("organizations",[]),
            "amounts": e.get("amounts",[])
        },
        "sentiment": str(result.get("sentiment","Neutral")).strip()
    }

HTML = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>DocuLens AI — Intelligent Document Analysis</title>
<link href="https://fonts.googleapis.com/css2?family=Syne:wght@400;600;700;800&family=DM+Mono:wght@300;400;500&display=swap" rel="stylesheet">
<style>
:root{--bg:#04060f;--surface:#0b0f1e;--surface2:#111827;--border:#1e2d47;--teal:#00e5c3;--teal-dim:#00b89e;--amber:#ffb830;--red:#ff4d6d;--green:#00e676;--text:#e8edf5;--muted:#5a6a85;--font-head:'Syne',sans-serif;--font-mono:'DM Mono',monospace}
*,*::before,*::after{box-sizing:border-box;margin:0;padding:0}
html{scroll-behavior:smooth}
body{background:var(--bg);color:var(--text);font-family:var(--font-mono);min-height:100vh;overflow-x:hidden}
body::before{content:'';position:fixed;inset:0;background-image:linear-gradient(rgba(0,229,195,0.03) 1px,transparent 1px),linear-gradient(90deg,rgba(0,229,195,0.03) 1px,transparent 1px);background-size:40px 40px;pointer-events:none;z-index:0}
body::after{content:'';position:fixed;width:600px;height:600px;border-radius:50%;background:radial-gradient(circle,rgba(0,229,195,0.06) 0%,transparent 70%);top:-200px;right:-200px;pointer-events:none;z-index:0}
.container{max-width:960px;margin:0 auto;padding:0 24px;position:relative;z-index:1}
header{padding:32px 0 24px;border-bottom:1px solid var(--border)}
.header-inner{display:flex;align-items:center;justify-content:space-between}
.logo{display:flex;align-items:center;gap:12px}
.logo-icon{width:40px;height:40px;background:var(--teal);clip-path:polygon(50% 0%,100% 25%,100% 75%,50% 100%,0% 75%,0% 25%);animation:pulse-glow 3s ease-in-out infinite}
@keyframes pulse-glow{0%,100%{filter:drop-shadow(0 0 8px rgba(0,229,195,0.6))}50%{filter:drop-shadow(0 0 20px rgba(0,229,195,1))}}
.logo-text{font-family:var(--font-head);font-size:22px;font-weight:800;letter-spacing:-0.5px}
.logo-text span{color:var(--teal)}
.badge{font-size:11px;padding:4px 10px;border:1px solid var(--teal);color:var(--teal);border-radius:2px;letter-spacing:2px;text-transform:uppercase}
.hero{padding:60px 0 40px;text-align:center}
.hero h1{font-family:var(--font-head);font-size:clamp(36px,6vw,64px);font-weight:800;line-height:1.05;letter-spacing:-2px;margin-bottom:16px}
.hero h1 .accent{color:var(--teal)}
.hero p{color:var(--muted);font-size:15px;max-width:480px;margin:0 auto 32px;line-height:1.7}
.format-pills{display:flex;gap:8px;justify-content:center;flex-wrap:wrap}
.pill{padding:6px 14px;background:var(--surface2);border:1px solid var(--border);border-radius:2px;font-size:12px;letter-spacing:1px;text-transform:uppercase;color:var(--muted);transition:all 0.2s}
.pill:hover{border-color:var(--teal);color:var(--teal)}
.upload-section{padding:40px 0}
.drop-zone{border:2px dashed var(--border);border-radius:4px;padding:60px 40px;text-align:center;cursor:pointer;transition:all 0.25s;background:var(--surface);position:relative;overflow:hidden}
.drop-zone::before{content:'';position:absolute;inset:0;background:linear-gradient(135deg,rgba(0,229,195,0.03),transparent);opacity:0;transition:opacity 0.3s}
.drop-zone:hover,.drop-zone.drag-over{border-color:var(--teal);background:rgba(0,229,195,0.04)}
.drop-zone:hover::before,.drop-zone.drag-over::before{opacity:1}
.drop-icon{font-size:48px;margin-bottom:16px;display:block;opacity:0.6}
.drop-zone h3{font-family:var(--font-head);font-size:20px;font-weight:700;margin-bottom:8px}
.drop-zone p{color:var(--muted);font-size:13px}
.drop-zone p span{color:var(--teal)}
input[type=file]{display:none}
.file-preview{display:none;margin-top:20px;padding:16px 20px;background:var(--surface2);border:1px solid var(--border);border-radius:4px;align-items:center;gap:12px}
.file-preview.show{display:flex}
.file-icon{width:40px;height:40px;background:rgba(0,229,195,0.1);border:1px solid rgba(0,229,195,0.3);border-radius:4px;display:flex;align-items:center;justify-content:center;font-size:18px;flex-shrink:0}
.file-info{flex:1;min-width:0}
.file-name{font-size:14px;font-weight:500;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}
.file-size{font-size:12px;color:var(--muted);margin-top:2px}
.remove-btn{background:none;border:none;color:var(--muted);cursor:pointer;font-size:18px;padding:4px;transition:color 0.2s}
.remove-btn:hover{color:var(--red)}
.api-row{margin-top:16px;display:flex;gap:12px;align-items:center}
.api-label{font-size:12px;color:var(--muted);white-space:nowrap;letter-spacing:1px;text-transform:uppercase}
.api-input{flex:1;background:var(--surface2);border:1px solid var(--border);border-radius:4px;padding:10px 14px;color:var(--text);font-family:var(--font-mono);font-size:13px;outline:none;transition:border-color 0.2s}
.api-input:focus{border-color:var(--teal)}
.analyze-btn{width:100%;margin-top:20px;padding:16px;background:var(--teal);color:#000;border:none;border-radius:4px;font-family:var(--font-head);font-size:16px;font-weight:700;letter-spacing:1px;text-transform:uppercase;cursor:pointer;transition:all 0.2s;position:relative;overflow:hidden}
.analyze-btn:hover{background:var(--teal-dim);transform:translateY(-1px);box-shadow:0 8px 32px rgba(0,229,195,0.3)}
.analyze-btn:disabled{opacity:0.5;cursor:not-allowed;transform:none;box-shadow:none}
.loading{display:none;margin-top:40px;text-align:center}
.loading.show{display:block}
.spinner{width:48px;height:48px;border:2px solid var(--border);border-top-color:var(--teal);border-radius:50%;animation:spin 0.8s linear infinite;margin:0 auto 16px}
@keyframes spin{to{transform:rotate(360deg)}}
.loading p{color:var(--muted);font-size:13px}
.loading-steps span{display:inline-block;animation:blink 1.4s ease-in-out infinite}
.loading-steps span:nth-child(2){animation-delay:0.2s}
.loading-steps span:nth-child(3){animation-delay:0.4s}
@keyframes blink{0%,80%,100%{opacity:0}40%{opacity:1}}
.results{display:none;margin-top:48px;animation:fadeUp 0.4s ease}
.results.show{display:block}
@keyframes fadeUp{from{opacity:0;transform:translateY(20px)}to{opacity:1;transform:translateY(0)}}
.results-header{display:flex;align-items:center;justify-content:space-between;margin-bottom:24px;padding-bottom:16px;border-bottom:1px solid var(--border)}
.results-title{font-family:var(--font-head);font-size:14px;font-weight:700;letter-spacing:2px;text-transform:uppercase;color:var(--muted)}
.results-file{font-size:13px;color:var(--teal)}
.sentiment-bar{display:flex;align-items:center;gap:16px;padding:20px 24px;background:var(--surface);border:1px solid var(--border);border-radius:4px;margin-bottom:20px}
.sentiment-label{font-size:12px;color:var(--muted);letter-spacing:2px;text-transform:uppercase}
.sentiment-badge{padding:6px 16px;border-radius:2px;font-family:var(--font-head);font-size:15px;font-weight:700;letter-spacing:1px}
.sentiment-badge.positive{background:rgba(0,230,118,0.15);color:var(--green);border:1px solid rgba(0,230,118,0.3)}
.sentiment-badge.neutral{background:rgba(255,184,48,0.15);color:var(--amber);border:1px solid rgba(255,184,48,0.3)}
.sentiment-badge.negative{background:rgba(255,77,109,0.15);color:var(--red);border:1px solid rgba(255,77,109,0.3)}
.summary-card{padding:24px;background:var(--surface);border:1px solid var(--border);border-left:3px solid var(--teal);border-radius:4px;margin-bottom:20px}
.card-label{font-size:11px;color:var(--teal);letter-spacing:2px;text-transform:uppercase;margin-bottom:12px;display:block}
.summary-text{font-size:14px;line-height:1.8;color:#c8d4e8}
.entities-grid{display:grid;grid-template-columns:repeat(2,1fr);gap:16px;margin-bottom:20px}
@media(max-width:600px){.entities-grid{grid-template-columns:1fr}}
.entity-card{padding:20px;background:var(--surface);border:1px solid var(--border);border-radius:4px;transition:border-color 0.2s}
.entity-card:hover{border-color:rgba(0,229,195,0.3)}
.entity-header{display:flex;align-items:center;gap:8px;margin-bottom:14px}
.entity-icon{font-size:16px}
.entity-type{font-size:11px;color:var(--muted);letter-spacing:2px;text-transform:uppercase}
.entity-count{margin-left:auto;font-size:11px;color:var(--teal);background:rgba(0,229,195,0.1);padding:2px 8px;border-radius:2px}
.entity-tags{display:flex;flex-wrap:wrap;gap:6px}
.entity-tag{padding:4px 10px;background:var(--surface2);border:1px solid var(--border);border-radius:2px;font-size:12px;color:var(--text)}
.entity-empty{font-size:12px;color:var(--muted);font-style:italic}
.json-section{margin-bottom:48px}
.json-header{display:flex;align-items:center;justify-content:space-between;margin-bottom:12px}
.copy-btn{background:none;border:1px solid var(--border);color:var(--muted);padding:6px 14px;border-radius:2px;font-family:var(--font-mono);font-size:12px;cursor:pointer;transition:all 0.2s}
.copy-btn:hover{border-color:var(--teal);color:var(--teal)}
.json-output{background:var(--surface);border:1px solid var(--border);border-radius:4px;padding:20px;font-size:12px;line-height:1.8;overflow-x:auto;white-space:pre;color:#7ec8a4;max-height:320px;overflow-y:auto}
.error-box{display:none;margin-top:20px;padding:16px 20px;background:rgba(255,77,109,0.08);border:1px solid rgba(255,77,109,0.3);border-radius:4px;color:var(--red);font-size:13px}
.error-box.show{display:block}
footer{border-top:1px solid var(--border);padding:24px 0;text-align:center;color:var(--muted);font-size:12px;margin-top:40px}
footer span{color:var(--teal)}
</style>
</head>
<body>
<div class="container">
  <header>
    <div class="header-inner">
      <div class="logo">
        <div class="logo-icon"></div>
        <div class="logo-text">Docu<span>Lens</span> AI</div>
      </div>
      <div class="badge">Track 2</div>
    </div>
  </header>

  <section class="hero">
    <h1>Intelligent<br><span class="accent">Document Analysis</span></h1>
    <p>Upload any document and extract structured insights — summaries, entities, and sentiment — powered by Gemini 1.5 Flash.</p>
    <div class="format-pills">
      <div class="pill">📄 PDF</div>
      <div class="pill">📝 DOCX</div>
      <div class="pill">🖼 PNG</div>
      <div class="pill">📷 JPEG</div>
      <div class="pill">🔍 OCR</div>
    </div>
  </section>

  <section class="upload-section">
    <div class="drop-zone" id="dropZone" onclick="document.getElementById('fileInput').click()">
      <span class="drop-icon">⬆</span>
      <h3>Drop your document here</h3>
      <p>or <span>browse files</span> &middot; PDF, DOCX, PNG, JPEG supported</p>
      <input type="file" id="fileInput" accept=".pdf,.docx,.doc,.png,.jpg,.jpeg,.tiff,.bmp">
    </div>

    <div class="file-preview" id="filePreview">
      <div class="file-icon" id="fileIcon">📄</div>
      <div class="file-info">
        <div class="file-name" id="fileName">—</div>
        <div class="file-size" id="fileSize">—</div>
      </div>
      <button class="remove-btn" onclick="removeFile()">✕</button>
    </div>

    <div class="api-row">
      <span class="api-label">x-api-key</span>
      <input class="api-input" type="text" id="apiKey" value="sk_track2_987654321" placeholder="Enter API key">
    </div>

    <button class="analyze-btn" id="analyzeBtn" onclick="analyze()">⚡ Analyze Document</button>
    <div class="error-box" id="errorBox"></div>

    <div class="loading" id="loading">
      <div class="spinner"></div>
      <p>Analyzing document<span class="loading-steps"><span>.</span><span>.</span><span>.</span></span></p>
    </div>
  </section>

  <section class="results" id="results">
    <div class="results-header">
      <span class="results-title">Analysis Results</span>
      <span class="results-file" id="resultsFile"></span>
    </div>
    <div class="sentiment-bar">
      <span class="sentiment-label">Overall Sentiment</span>
      <span class="sentiment-badge" id="sentimentBadge">—</span>
    </div>
    <div class="summary-card">
      <span class="card-label">Summary</span>
      <p class="summary-text" id="summaryText">—</p>
    </div>
    <div class="entities-grid">
      <div class="entity-card">
        <div class="entity-header"><span class="entity-icon">👤</span><span class="entity-type">Names</span><span class="entity-count" id="namesCount">0</span></div>
        <div class="entity-tags" id="namesTags"></div>
      </div>
      <div class="entity-card">
        <div class="entity-header"><span class="entity-icon">📅</span><span class="entity-type">Dates</span><span class="entity-count" id="datesCount">0</span></div>
        <div class="entity-tags" id="datesTags"></div>
      </div>
      <div class="entity-card">
        <div class="entity-header"><span class="entity-icon">🏢</span><span class="entity-type">Organizations</span><span class="entity-count" id="orgsCount">0</span></div>
        <div class="entity-tags" id="orgsTags"></div>
      </div>
      <div class="entity-card">
        <div class="entity-header"><span class="entity-icon">💰</span><span class="entity-type">Amounts</span><span class="entity-count" id="amountsCount">0</span></div>
        <div class="entity-tags" id="amountsTags"></div>
      </div>
    </div>
    <div class="json-section">
      <div class="json-header">
        <span class="card-label" style="margin:0">Raw JSON Response</span>
        <button class="copy-btn" onclick="copyJson()">Copy JSON</button>
      </div>
      <div class="json-output" id="jsonOutput"></div>
    </div>
  </section>

  <footer>Built with <span>Gemini 1.5 Flash</span> &middot; FastAPI &middot; Tesseract OCR &middot; Track 2 Submission</footer>
</div>

<script>
let selectedFile=null;
const fileIcons={pdf:'📄',docx:'📝',doc:'📝',png:'🖼',jpg:'📷',jpeg:'📷',tiff:'🖼',bmp:'🖼'};
function formatBytes(b){if(b<1024)return b+' B';if(b<1048576)return(b/1024).toFixed(1)+' KB';return(b/1048576).toFixed(1)+' MB';}
function getExt(name){return name.split('.').pop().toLowerCase();}
function setFile(file){
  selectedFile=file;
  const ext=getExt(file.name);
  document.getElementById('fileName').textContent=file.name;
  document.getElementById('fileSize').textContent=formatBytes(file.size);
  document.getElementById('fileIcon').textContent=fileIcons[ext]||'📄';
  document.getElementById('filePreview').classList.add('show');
  document.getElementById('errorBox').classList.remove('show');
  document.getElementById('results').classList.remove('show');
}
function removeFile(){
  selectedFile=null;
  document.getElementById('fileInput').value='';
  document.getElementById('filePreview').classList.remove('show');
}
document.getElementById('fileInput').addEventListener('change',e=>{if(e.target.files[0])setFile(e.target.files[0]);});
const dz=document.getElementById('dropZone');
dz.addEventListener('dragover',e=>{e.preventDefault();dz.classList.add('drag-over');});
dz.addEventListener('dragleave',()=>dz.classList.remove('drag-over'));
dz.addEventListener('drop',e=>{e.preventDefault();dz.classList.remove('drag-over');if(e.dataTransfer.files[0])setFile(e.dataTransfer.files[0]);});
function showError(msg){const b=document.getElementById('errorBox');b.textContent='⚠ '+msg;b.classList.add('show');}
function renderTags(cid,countId,items){
  const c=document.getElementById(cid);
  document.getElementById(countId).textContent=items.length;
  c.innerHTML=items.length?items.map(i=>`<span class="entity-tag">${i}</span>`).join(''):'<span class="entity-empty">None found</span>';
}
async function analyze(){
  if(!selectedFile){showError('Please select a document first.');return;}
  const apiKey=document.getElementById('apiKey').value.trim();
  if(!apiKey){showError('Please enter your API key.');return;}
  const ext=getExt(selectedFile.name);
  document.getElementById('analyzeBtn').disabled=true;
  document.getElementById('loading').classList.add('show');
  document.getElementById('results').classList.remove('show');
  document.getElementById('errorBox').classList.remove('show');
  try{
    const ab=await selectedFile.arrayBuffer();
    const bytes=new Uint8Array(ab);let b64='';for(let i=0;i<bytes.length;i+=8192){b64+=String.fromCharCode(...bytes.subarray(i,i+8192));}b64=btoa(b64);
    const resp=await fetch('/api/document-analyze',{
      method:'POST',
      headers:{'Content-Type':'application/json','x-api-key':apiKey},
      body:JSON.stringify({fileName:selectedFile.name,fileType:ext,fileBase64:b64})
    });
    const data=await resp.json();
    if(!resp.ok){showError(data.detail||'Analysis failed.');return;}
    document.getElementById('resultsFile').textContent=data.fileName;
    document.getElementById('summaryText').textContent=data.summary;
    const s=data.sentiment.toLowerCase();
    const sb=document.getElementById('sentimentBadge');
    sb.textContent=data.sentiment;
    sb.className='sentiment-badge '+(s==='positive'?'positive':s==='negative'?'negative':'neutral');
    renderTags('namesTags','namesCount',data.entities.names||[]);
    renderTags('datesTags','datesCount',data.entities.dates||[]);
    renderTags('orgsTags','orgsCount',data.entities.organizations||[]);
    renderTags('amountsTags','amountsCount',data.entities.amounts||[]);
    document.getElementById('jsonOutput').textContent=JSON.stringify(data,null,2);
    document.getElementById('results').classList.add('show');
  }catch(err){showError('Network error: '+err.message);}
  finally{
    document.getElementById('analyzeBtn').disabled=false;
    document.getElementById('loading').classList.remove('show');
  }
}
function copyJson(){
  navigator.clipboard.writeText(document.getElementById('jsonOutput').textContent).then(()=>{
    const b=document.querySelector('.copy-btn');b.textContent='Copied!';setTimeout(()=>b.textContent='Copy JSON',2000);
  });
}
</script>
</body>
</html>"""

@app.get("/", response_class=HTMLResponse)
def frontend():
    return HTML

@app.get("/health")
def health():
    return {"status": "healthy"}

@app.post("/api/document-analyze", response_model=DocumentResponse)
def document_analyze(request: DocumentRequest, x_api_key: str = Header(None)):
    if not x_api_key or x_api_key != API_KEY:
        raise HTTPException(status_code=401, detail="Unauthorized: Invalid or missing x-api-key header")
    try:
        file_bytes = base64.b64decode(request.fileBase64)
    except:
        raise HTTPException(status_code=400, detail="Invalid base64 string")
    if not file_bytes:
        raise HTTPException(status_code=400, detail="Empty file")
    try:
        text = get_text(request.fileType, file_bytes)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=422, detail=f"Extraction failed: {str(e)}")
    if not text or len(text.strip()) < 5:
        raise HTTPException(status_code=422, detail="No readable text found in document")
    try:
        result = analyze(text, request.fileName)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI analysis failed: {str(e)}")
    return DocumentResponse(status="success", fileName=request.fileName,
        summary=result["summary"], entities=result["entities"], sentiment=result["sentiment"])
